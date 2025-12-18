#!/usr/bin/env python3
"""
File Converter - Convert between text document formats
Supports: txt, md, html, docx, pdf

For better conversion quality, consider using native tools:
  - pandoc: Universal document converter
  - pdftotext: PDF to text (poppler-utils package)
"""

import argparse
import re
import sys
from pathlib import Path

# Check for optional dependencies
DOCX_AVAILABLE = False
MARKDOWNIFY_AVAILABLE = False
PDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    from markdownify import markdownify

    MARKDOWNIFY_AVAILABLE = True
except ImportError:
    pass

try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    pass


def txt_to_md(text: str, title: str = None) -> str:
    """
    Convert plain text to markdown with basic formatting detection.

    Args:
        text: Plain text content
        title: Optional title for the document

    Returns:
        Markdown formatted text
    """
    lines = text.split("\n")
    result = []

    if title:
        result.append(f"# {title}\n")

    for line in lines:
        line = line.rstrip()

        # Detect potential headers (all caps, 3+ words)
        if line.isupper() and len(line.split()) >= 3:
            result.append(f"## {line.title()}")
        # Lines ending with colon might be subheaders
        elif line.endswith(":") and len(line) < 60:
            result.append(f"### {line}")
        else:
            result.append(line)

    return "\n".join(result)


def html_to_md(html: str) -> str:
    """
    Convert HTML to markdown.

    Args:
        html: HTML content

    Returns:
        Markdown formatted text

    Raises:
        SystemExit: If markdownify is not installed
    """
    if not MARKDOWNIFY_AVAILABLE:
        print("Error: markdownify not installed", file=sys.stderr)
        print("Install with: pip install markdownify", file=sys.stderr)
        sys.exit(1)

    return markdownify(html, heading_style="ATX", strip=["script", "style"])


def md_to_txt(markdown: str) -> str:
    """
    Convert markdown to plain text by stripping formatting.

    Args:
        markdown: Markdown content

    Returns:
        Plain text
    """
    text = markdown

    # Remove headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)

    # Remove bold
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)

    # Remove italic
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)

    # Remove links but keep text
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)

    # Remove images
    text = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"\1", text)

    # Remove code blocks
    text = re.sub(r"```[\s\S]*?```", "", text)

    # Remove inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # Remove blockquotes
    text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)

    # Remove horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)

    # Clean up excessive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def docx_to_md(path: str) -> str:
    """
    Convert DOCX to markdown.

    Args:
        path: Path to DOCX file

    Returns:
        Markdown formatted text

    Raises:
        SystemExit: If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style = para.style.name.lower() if para.style else ""

        # Convert styles to markdown
        if "heading 1" in style or "title" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        elif "heading" in style:
            lines.append(f"#### {text}")
        else:
            # Handle bold/italic formatting
            formatted_text = []
            for run in para.runs:
                run_text = run.text
                if run.bold and run.italic:
                    formatted_text.append(f"***{run_text}***")
                elif run.bold:
                    formatted_text.append(f"**{run_text}**")
                elif run.italic:
                    formatted_text.append(f"*{run_text}*")
                else:
                    formatted_text.append(run_text)
            lines.append("".join(formatted_text))

    return "\n\n".join(lines)


def docx_to_txt(path: str) -> str:
    """
    Convert DOCX to plain text.

    Args:
        path: Path to DOCX file

    Returns:
        Plain text

    Raises:
        SystemExit: If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def md_to_docx(markdown: str, output_path: str) -> None:
    """
    Convert markdown to DOCX.

    Args:
        markdown: Markdown content
        output_path: Path to save DOCX file

    Raises:
        SystemExit: If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Handle headers
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = min(len(m.group(1)), 9)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Handle code blocks
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                doc.add_paragraph("\n".join(code_lines)).style = "Quote"
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            # Strip markdown formatting
            line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
            line = re.sub(r"\*([^*]+)\*", r"\1", line)
            line = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", line)
            doc.add_paragraph(line)

        i += 1

    doc.save(output_path)


def pdf_to_txt(path: str) -> str:
    """
    Convert PDF to plain text using PyPDF2.

    Note: For better quality, use system tools:
      pdftotext input.pdf output.txt

    Args:
        path: Path to PDF file

    Returns:
        Extracted text

    Raises:
        SystemExit: If PyPDF2 is not installed
    """
    if not PDF_AVAILABLE:
        print("Error: PyPDF2 not installed", file=sys.stderr)
        print("Install with: pip install PyPDF2", file=sys.stderr)
        print("\nAlternatively, use system tools for better quality:", file=sys.stderr)
        print("  sudo dnf install poppler-utils", file=sys.stderr)
        print("  pdftotext input.pdf output.txt", file=sys.stderr)
        sys.exit(1)

    text_parts = []

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text = page.extract_text() or ""
            text_parts.append(text)

    return "\n\n".join(text_parts).strip()


def pdf_to_md(path: str) -> str:
    """
    Convert PDF to markdown.

    Args:
        path: Path to PDF file

    Returns:
        Markdown formatted text
    """
    title = Path(path).stem
    return txt_to_md(pdf_to_txt(path), title)


def get_format(path: str) -> str:
    """
    Detect file format from extension.

    Args:
        path: File path

    Returns:
        Format identifier (txt, md, html, docx, pdf)
    """
    ext = Path(path).suffix.lower()
    format_map = {
        ".txt": "txt",
        ".md": "md",
        ".markdown": "md",
        ".html": "html",
        ".htm": "html",
        ".docx": "docx",
        ".pdf": "pdf",
    }
    return format_map.get(ext, "txt")


def convert_file(input_path: str, output_format: str) -> tuple:
    """
    Convert a file to the specified format.

    Args:
        input_path: Path to input file
        output_format: Target format (txt, md, html, docx)

    Returns:
        Tuple of (content, extension) or (None, extension) for docx output

    Raises:
        ValueError: If conversion is not supported
        IOError: If file cannot be read
    """
    input_format = get_format(input_path)

    # No conversion needed
    if input_format == output_format:
        with open(input_path, "r", encoding="utf-8") as f:
            return f.read(), f".{output_format}"

    # Define conversion functions
    conversions = {
        ("txt", "md"): lambda: (
            txt_to_md(
                open(input_path, "r", encoding="utf-8").read(), Path(input_path).stem
            ),
            ".md",
        ),
        ("md", "txt"): lambda: (
            md_to_txt(open(input_path, "r", encoding="utf-8").read()),
            ".txt",
        ),
        ("html", "md"): lambda: (
            html_to_md(open(input_path, "r", encoding="utf-8").read()),
            ".md",
        ),
        ("html", "txt"): lambda: (
            md_to_txt(html_to_md(open(input_path, "r", encoding="utf-8").read())),
            ".txt",
        ),
        ("docx", "md"): lambda: (docx_to_md(input_path), ".md"),
        ("docx", "txt"): lambda: (docx_to_txt(input_path), ".txt"),
        ("md", "docx"): lambda: (None, ".docx"),  # Special case - writes directly
        ("pdf", "txt"): lambda: (pdf_to_txt(input_path), ".txt"),
        ("pdf", "md"): lambda: (pdf_to_md(input_path), ".md"),
    }

    conversion_key = (input_format, output_format)
    if conversion_key in conversions:
        return conversions[conversion_key]()

    raise ValueError(f"Conversion from {input_format} to {output_format} not supported")


def main():
    """Main entry point for CLI usage."""
    parser = argparse.ArgumentParser(
        description="Convert between document formats (txt, md, html, docx, pdf)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s document.pdf --to txt
  %(prog)s file.docx --to md -o output.md
  %(prog)s *.txt --to md --batch
  %(prog)s file1.pdf file2.pdf --to txt --batch -o converted/

Native Linux Alternatives (often better quality):
  # Universal converter (highly recommended)
  pandoc input.docx -o output.md
  pandoc input.md -o output.pdf

  # PDF to text (better than PyPDF2)
  pdftotext input.pdf output.txt

  # Install tools
  sudo dnf install pandoc poppler-utils
        """,
    )

    parser.add_argument("input_files", nargs="+", help="Input file(s) to convert")

    parser.add_argument(
        "--to",
        required=True,
        choices=["txt", "md", "html", "docx"],
        help="Target format",
    )

    parser.add_argument(
        "-o", "--output", help="Output file or directory (for batch mode)"
    )

    parser.add_argument(
        "--batch",
        action="store_true",
        help="Batch mode - convert multiple files to directory",
    )

    args = parser.parse_args()

    try:
        # Batch mode or multiple files
        if args.batch or len(args.input_files) > 1:
            output_dir = Path(args.output) if args.output else Path("converted")
            output_dir.mkdir(parents=True, exist_ok=True)

            success_count = 0
            total = len(args.input_files)

            print(f"Converting {total} file(s) to {args.to.upper()}...")
            print(f"Output directory: {output_dir}\n")

            for input_file in args.input_files:
                input_path = Path(input_file)

                if not input_path.exists():
                    print(
                        f"  ⚠ Warning: {input_file} not found, skipping",
                        file=sys.stderr,
                    )
                    continue

                try:
                    content, ext = convert_file(str(input_path), args.to)
                    output_file = output_dir / f"{input_path.stem}{ext}"

                    if args.to == "docx" and content is None:
                        # Special handling for markdown to docx
                        with open(input_path, "r", encoding="utf-8") as f:
                            md_to_docx(f.read(), str(output_file))
                    else:
                        with open(output_file, "w", encoding="utf-8") as f:
                            f.write(content)

                    print(f"  ✓ {input_file} -> {output_file}")
                    success_count += 1

                except Exception as e:
                    print(f"  ✗ Error converting {input_file}: {e}", file=sys.stderr)

            print(f"\n{'='*60}")
            print(f"Completed: {success_count}/{total} files converted successfully")
            print(f"{'='*60}")

        # Single file mode
        else:
            input_file = args.input_files[0]
            input_path = Path(input_file)

            if not input_path.exists():
                print(f"Error: {input_file} not found", file=sys.stderr)
                sys.exit(1)

            content, ext = convert_file(input_file, args.to)
            output_file = args.output if args.output else f"{input_path.stem}{ext}"

            if args.to == "docx" and content is None:
                # Special handling for markdown to docx
                with open(input_file, "r", encoding="utf-8") as f:
                    md_to_docx(f.read(), output_file)
            else:
                with open(output_file, "w", encoding="utf-8") as f:
                    f.write(content)

            print(f"✓ Converted: {input_file} -> {output_file}")

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
