#!/usr/bin/env python3
"""
Unit tests for file_converter.py.
"""

import os
import sys
from pathlib import Path

import pytest

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from file_converter import (
    DOCX_AVAILABLE,
    MARKDOWNIFY_AVAILABLE,
    PDF_AVAILABLE,
    convert_file,
    docx_to_md,
    docx_to_txt,
    get_format,
    html_to_md,
    md_to_docx,
    md_to_txt,
    pdf_to_md,
    pdf_to_txt,
    txt_to_md,
)

# --- Fixtures ---


@pytest.fixture
def sample_txt():
    return "This is a sample text."


@pytest.fixture
def sample_md():
    return "# This is a sample markdown."


@pytest.fixture
def sample_html():
    return "<h1>This is a sample html.</h1>"


# --- Tests ---


def test_get_format():
    assert get_format("test.txt") == "txt"
    assert get_format("test.md") == "md"
    assert get_format("test.html") == "html"
    assert get_format("test.docx") == "docx"
    assert get_format("test.pdf") == "pdf"
    assert get_format("test.unknown") == "txt"


def test_txt_to_md(sample_txt):
    md = txt_to_md(sample_txt)
    assert "This is a sample text." in md


def test_md_to_txt(sample_md):
    txt = md_to_txt(sample_md)
    assert "This is a sample markdown." in txt


@pytest.mark.skipif(not MARKDOWNIFY_AVAILABLE, reason="markdownify not installed")
def test_html_to_md(sample_html):
    md = html_to_md(sample_html)
    assert "# This is a sample html." in md


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_to_txt():
    # Create a dummy docx file
    from docx import Document

    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.save("test.docx")

    txt = docx_to_txt("test.docx")
    assert "This is a test document." in txt
    os.remove("test.docx")


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_to_md():
    # Create a dummy docx file
    from docx import Document

    doc = Document()
    doc.add_heading("This is a heading.", level=1)
    doc.save("test.docx")

    md = docx_to_md("test.docx")
    assert "# This is a heading." in md
    os.remove("test.docx")


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_md_to_docx(sample_md):
    md_to_docx(sample_md, "test.docx")
    assert os.path.exists("test.docx")
    os.remove("test.docx")


@pytest.mark.skipif(not PDF_AVAILABLE, reason="PyPDF2 not installed")
def test_pdf_to_txt():
    # This is a bit tricky to test without a real PDF file.
    # We will just check if the function runs without errors.
    # Create a dummy pdf file
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open("test.pdf", "wb") as f:
        writer.write(f)

    txt = pdf_to_txt("test.pdf")
    assert isinstance(txt, str)
    os.remove("test.pdf")


@pytest.mark.skipif(not PDF_AVAILABLE, reason="PyPDF2 not installed")
def test_pdf_to_md():
    # This is a bit tricky to test without a real PDF file.
    # We will just check if the function runs without errors.
    # Create a dummy pdf file
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open("test.pdf", "wb") as f:
        writer.write(f)

    md = pdf_to_md("test.pdf")
    assert isinstance(md, str)
    os.remove("test.pdf")


def test_convert_file(sample_txt):
    with open("test.txt", "w") as f:
        f.write(sample_txt)

    content, ext = convert_file("test.txt", "md")
    assert ext == ".md"
    assert "This is a sample text." in content
    os.remove("test.txt")
